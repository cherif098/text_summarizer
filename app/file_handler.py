import docx2txt
import PyPDF2
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from io import BytesIO
from pathlib import Path
import os
import logging
from textwrap import wrap

def read_word(file):
    """
    Extract text from a Word document (.docx).
    """
    try:
        # Vérification si le fichier est un objet de type 'file-like'
        if hasattr(file, 'read'):
            # Si c'est un objet comme BytesIO (ex : fichier téléchargé)
            file_like = BytesIO(file.read())
            doc = Document(file_like)
        elif isinstance(file, str):  # Si c'est un chemin de fichier
            doc = Document(file)
        else:
            raise ValueError("Invalid file input. Expected a file path or a file-like object.")
        
        # Extraction du texte du document Word
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        logging.error(f"Error reading Word file: {e}")
        return ""


def read_pdf(file):
    """
    Extract text from a PDF file.
    """
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text
        return text
    except Exception as e:
        print(f"Error reading PDF file: {e}")
        return ""


def save_as_word(text):
    try:
        if not text.strip():
            raise ValueError("The input text is empty.")
        output = BytesIO()
        doc = Document()
        doc.add_paragraph(text)
        doc.save(output)
        output.seek(0)
        return output
    except Exception as e:
        logging.error(f"Error saving Word file: {e}")
        return None

def save_as_pdf(text):
    """
    Save text as a PDF, handling line wrapping, and return the PDF as a BytesIO object.
    """
    try:
        if isinstance(text, list):
            text = " ".join(map(str, text))  # Convert list to a single string
        elif not isinstance(text, str):
            raise ValueError("The input text must be a string or a list of strings.")
        if not text.strip():
            raise ValueError("The input text is empty.")

        output = BytesIO()
        c = canvas.Canvas(output, pagesize=letter)
        width, height = letter

        # Set font and margins
        c.setFont("Helvetica", 12)
        x_margin, y_margin = 72, height - 72
        line_height = 14
        max_width = width - 2 * x_margin

        # Wrap text to fit within the page width
        from textwrap import wrap  # Ensure the correct import is present
        wrapped_text = wrap(text, width=int(max_width / 7))

        # Draw each line
        for line in wrapped_text:
            if y_margin < 72:  # If space runs out, create a new page
                c.showPage()
                y_margin = height - 72
                c.setFont("Helvetica", 12)
            c.drawString(x_margin, y_margin, line)
            y_margin -= line_height
        
        c.showPage()
        c.save()
        output.seek(0)
        return output
    except Exception as e:
        logging.error(f"Error creating PDF file: {e}")
        return None
